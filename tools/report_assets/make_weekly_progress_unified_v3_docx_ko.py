"""Unified weekly progress report v3 (Korean) for 2026-04-10.

Cleaner structure: fewer tables, figures inline with text, no file-path lists.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
WEEKLY_DIR = ROOT / "outputs" / "report_progress_explainer" / "weekly_progress_20260410"
COMPARE_DIR = ROOT / "outputs" / "report_progress_explainer" / "trot_stock_vs_custom_same_scenarios_20260410"
DOCX_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410_v3.docx"

SAME = ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios"
CRAWL_RELAX = ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck"

RUNS = {
    "same_straight_stock": SAME / "stock_straight_20s" / "episode_000" / "summary.json",
    "same_straight_custom": SAME / "custom_straight_20s" / "episode_000" / "summary.json",
    "same_turn_stock": SAME / "stock_turn_10s" / "episode_000" / "summary.json",
    "same_turn_custom": SAME / "custom_turn_10s" / "episode_000" / "summary.json",
    "same_disturb_stock": SAME / "stock_disturb_4s" / "episode_000" / "summary.json",
    "same_disturb_custom": SAME / "custom_disturb_4s" / "episode_000" / "summary.json",
    "before_straight": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_straight_default_20s" / "episode_000" / "summary.json",
    "before_turn": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_turn_default_10s" / "episode_000" / "summary.json",
    "before_disturb": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_baseline_before" / "episode_000" / "summary.json",
    "crawl_baseline": CRAWL_RELAX / "baseline_fy015_grf035" / "episode_000" / "summary.json",
    "crawl_fy100_grf035": CRAWL_RELAX / "fy100_grf035" / "episode_000" / "summary.json",
    "crawl_fy015_grf100": CRAWL_RELAX / "fy015_grf100" / "episode_000" / "summary.json",
    "crawl_fy100_grf100": CRAWL_RELAX / "fy100_grf100" / "episode_000" / "summary.json",
}

FIGURES = {
    "trot_fyscale100": WEEKLY_DIR / "weekly_trot_fyscale100_recheck.png",
    "trot_failure": WEEKLY_DIR / "weekly_failure_ablation_trot.png",
    "crawl_force_relax": WEEKLY_DIR / "weekly_crawl_force_relax_recheck.png",
    "same_straight": COMPARE_DIR / "trot_straight_stock_vs_custom.png",
    "same_turn": COMPARE_DIR / "trot_turn_stock_vs_custom.png",
    "same_disturb": COMPARE_DIR / "trot_disturbance_stock_vs_custom.png",
}


def _load(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def _fmt(v: float, d: int = 3) -> str:
    return f"{v:.{d}f}"


def _set_style(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Malgun Gothic"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    s.font.size = Pt(10)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = val
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)


def _add_fig(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.exists():
        doc.add_paragraph(f"(누락: {path.name})")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)


def build() -> None:
    data = {k: _load(v) for k, v in RUNS.items()}
    doc = Document()
    _set_style(doc)

    # ── Title ──
    title = doc.add_heading("이번 주 진행 보고", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Quadruped-PyMPC MuJoCo integration + custom linear_osqp backend")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ══════════════════════════════════════════════════════════
    # Section 1: 한 줄 요약 + 핵심 표
    # ══════════════════════════════════════════════════════════
    doc.add_heading("1. 핵심 결과", level=1)
    doc.add_paragraph(
        "custom linear_osqp 경로의 trot이 이번 주에 크게 개선되었습니다. "
        "아래 표는 같은 시나리오 조건(aliengo, flat, trot)에서 stock sampling과 custom linear_osqp를 "
        "다시 돌려서 비교한 결과입니다."
    )

    main_rows = []
    for label, scenario, metric in [
        ("Straight 20 s |roll|", "straight", "mean_abs_roll"),
        ("Straight 20 s |pitch|", "straight", "mean_abs_pitch"),
        ("Turn 10 s |roll|", "turn", "mean_abs_roll"),
        ("Turn 10 s |pitch|", "turn", "mean_abs_pitch"),
        ("Disturb 4 s |roll|", "disturb", "mean_abs_roll"),
        ("Disturb 4 s |pitch|", "disturb", "mean_abs_pitch"),
    ]:
        before_key = f"before_{scenario}"
        stock_key = f"same_{scenario}_stock"
        custom_key = f"same_{scenario}_custom"
        bv = data.get(before_key, {}).get(metric)
        sv = data.get(stock_key, {}).get(metric)
        cv = data.get(custom_key, {}).get(metric)
        main_rows.append([
            label,
            _fmt(bv) if bv is not None else "-",
            _fmt(cv) if cv is not None else "-",
            _fmt(sv) if sv is not None else "-",
        ])
    _add_table(doc, ["Scenario", "이번 주 전", "이번 주 후", "Stock"], main_rows)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Roll/pitch 기준으로 전 시나리오에서 custom이 stock 이하입니다. "
        "다만 turn의 forward velocity tracking(mean_vx)은 stock이 더 높으므로, "
        "posture quality 개선이지 전체 성능 우위라고 일반화하지는 않습니다."
    )

    # ══════════════════════════════════════════════════════════
    # Section 2: 핵심 발견 + 과정
    # ══════════════════════════════════════════════════════════
    doc.add_heading("2. 핵심 발견: post-solve lateral force 축소 구조가 trot turn의 병목이었다", level=1)

    doc.add_paragraph(
        "현재 제가 구현한 Python linear_osqp integration 경로에는, "
        "QP가 계산한 lateral force(fy)를 post-solve 단계에서 80% 축소하는 구조가 있었습니다 "
        "(fy_scale=0.20). 이는 초기 integration 단계에서 full lateral force가 불안정했을 때 "
        "도입한 보수적 안정화 선택이었습니다."
    )
    doc.add_paragraph(
        "이번 주에 이 구조가 더 이상 필요하지 않다는 것을 확인하고 제거(fy_scale=1.0)했습니다. "
        "그 과정은 다음과 같습니다."
    )

    steps = [
        ("Q weighting 강화 시도",
         "Q_theta_roll을 pitch와 대칭으로 올렸지만 turn roll은 변하지 않았습니다. "
         "\"왜 Q를 올려도 효과가 없는가?\"가 다음 질문이 되었습니다."),
        ("원인 추적",
         "QP가 풀어낸 fy를 post-solve에서 fy_scale=0.20으로 축소하는 구조를 확인했습니다. "
         "Q를 올려도 그 효과가 실제 적용 force까지 충분히 전달되지 않았을 가능성이 큽니다."),
        ("부분 개선: 축소 비율을 roll에 따라 조절",
         "dynamic_fy_roll_gain=0.25를 추가해 roll이 클 때 fy를 덜 깎도록 했습니다. "
         "turn roll이 0.024에서 0.018로 부분 개선되었습니다."),
        ("완전 제거: fy_scale=1.0",
         "축소 자체를 제거했습니다. turn roll 0.024 -> 0.011 (stock 0.014 이하), "
         "전 시나리오에서 안정적. 초기 안정화 구조가 더 이상 필요하지 않았습니다."),
    ]
    for i, (title_text, body) in enumerate(steps, 1):
        p = doc.add_paragraph()
        runner = p.add_run(f"{i}. {title_text}: ")
        runner.bold = True
        runner.font.size = Pt(10)
        p.add_run(body).font.size = Pt(10)

    doc.add_paragraph("")
    _add_fig(doc, FIGURES["trot_failure"],
             "그림 1. Q weighting 강화(실패)와 fy_scale 조정(효과 있었던 방향) 비교")

    doc.add_paragraph("")
    _add_fig(doc, FIGURES["trot_fyscale100"],
             "그림 2. fy_scale=1.0 최종 결과 -- 전 시나리오에서 stock 이하")

    # pitch bias
    doc.add_paragraph("")
    doc.add_paragraph(
        "별도로, pitch는 noise가 아니라 일관된 양의 방향 bias였습니다 "
        "(mean signed pitch = mean |pitch|). "
        "pitch_ref_offset을 -0.01에서 -0.03으로 조정해 상쇄했습니다."
    )

    # ══════════════════════════════════════════════════════════
    # Section 3: trot 시나리오별 비교 (그래프 중심)
    # ══════════════════════════════════════════════════════════
    doc.add_heading("3. 시나리오별 stock vs custom 비교", level=1)
    doc.add_paragraph(
        "아래 그래프는 같은 조건(speed, yaw rate, disturbance pulse)에서 "
        "stock sampling과 custom linear_osqp를 나란히 비교한 것입니다. "
        "height, roll/pitch time-series, foot-contact timeline을 함께 보여줍니다."
    )

    _add_fig(doc, FIGURES["same_straight"],
             "그림 3. Straight 20 s: stock(좌) vs custom(우)")
    _add_fig(doc, FIGURES["same_turn"],
             "그림 4. Turn 10 s: stock(좌) vs custom(우)")
    _add_fig(doc, FIGURES["same_disturb"],
             "그림 5. Disturbance 4 s: stock(좌) vs custom(우)")

    # ══════════════════════════════════════════════════════════
    # Section 4: crawl
    # ══════════════════════════════════════════════════════════
    doc.add_heading("4. crawl에는 같은 접근이 통하지 않았다", level=1)
    doc.add_paragraph(
        "trot에서 효과가 컸던 force limit 완화를 crawl에도 적용해 "
        "fy_scale / grf_max_scale 4개 조합을 테스트했지만, "
        "전부 baseline(13.540 s)보다 나빠졌습니다."
    )

    crawl_rows = [
        ["baseline (fy=0.15, grf=0.35)", "13.540", "0.158", "0.038"],
        ["fy=1.0, grf=0.35", "7.190", "0.172", "0.056"],
        ["fy=0.15, grf=1.0", "9.568", "0.080", "0.028"],
        ["fy=1.0, grf=1.0", "8.390", "0.068", "0.031"],
    ]
    _add_table(doc, ["설정", "duration [s]", "mean|roll|", "mean|pitch|"], crawl_rows)

    doc.add_paragraph("")
    _add_fig(doc, FIGURES["crawl_force_relax"],
             "그림 6. crawl force-relax 재검증 -- 4개 조합 모두 baseline보다 악화")

    doc.add_paragraph("")
    doc.add_paragraph(
        "crawl은 upstream Quadruped-PyMPC README에서 공식 지원 gait으로 명시되어 있지 않습니다 "
        "(trot, pace, bound만 문서화). 같은 환경에서 stock sampling도 약 1.4-1.7 s에서 종료됩니다. "
        "현재 custom 경로의 13.5 s는 stock 대비 연장이지만, body height가 크게 낮은 상태에서의 "
        "생존이므로 functional locomotion이 아닌 contact-transition diagnostic으로 해석하는 것이 적절합니다."
    )

    # ══════════════════════════════════════════════════════════
    # Section 5: 시사점 + 남은 문제
    # ══════════════════════════════════════════════════════════
    doc.add_heading("5. 시사하는 점과 남은 문제", level=1)
    doc.add_paragraph(
        "이번 주 결과는 \"MPC cost를 더 키우는 것\"보다 "
        "\"이미 풀어낸 답이 실제로 어떻게 적용되는지\"를 점검하는 것이 "
        "더 중요할 수 있다는 점을 보여줍니다. "
        "trot에서는 post-solve scaling이 핵심 병목이었고, "
        "crawl에서는 같은 방식이 통하지 않아 force authority가 아닌 "
        "contact transition seam 쪽의 문제가 더 크다는 해석이 강화되었습니다."
    )
    doc.add_paragraph("남은 문제:")
    remaining = [
        "crawl late-seam failure (13.5 s): force가 아닌 transition timing / stabilization 쪽 조사 필요",
        "fy_scale=1.0에서 dynamic_fy_roll_gain=0.25가 여전히 필요한지 재검증",
        "turn forward velocity tracking: posture는 개선됐지만 mean_vx는 stock이 더 높음",
        "장기 안정성: 20 s 이상, 더 강한 disturbance 등에서 추가 검증",
    ]
    for item in remaining:
        doc.add_paragraph(item, style="List Bullet")

    # ══════════════════════════════════════════════════════════
    # 부록
    # ══════════════════════════════════════════════════════════
    doc.add_heading("부록: 시나리오 조건", level=1)
    cond_rows = [
        ["straight", "20 s", "0.12 m/s", "0.0 rad/s", "없음"],
        ["turn", "10 s", "0.10 m/s", "0.3 rad/s", "없음"],
        ["disturbance", "4 s", "0.12 m/s", "0.0 rad/s", "x:0.5:0.25:4.0, x:2.3:0.25:8.0"],
    ]
    _add_table(doc, ["scenario", "duration", "speed", "yaw rate", "disturbance"], cond_rows)
    doc.add_paragraph(
        "공통: aliengo, flat ground, trot gait. "
        "이번 주 전(before)은 fy_scale=0.20, dynamic_fy_roll_gain=0.0, pitch_ref_offset=-0.01. "
        "이번 주 후(after)는 fy_scale=1.0, dynamic_fy_roll_gain=0.25, pitch_ref_offset=-0.03."
    )

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    build()
