from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "outputs" / "report_progress_explainer"
CRAWL_DIR = OUT_DIR / "crawl_same_scenario"
DOCX_PATH = OUT_DIR / "quadruped_mpc_progress_update.docx"


def set_normal_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True


def add_picture(doc: Document, path: Path, width: float, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption(doc, caption_text)


def build_doc() -> Path:
    doc = Document()
    set_normal_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_page_number(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quadruped MPC Progress Update")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Python baseline -> stock Quadruped-PyMPC integration -> crawl diagnostic comparison")

    paragraph(
        doc,
        "This draft summarizes the progress up to this week. The main goal has been to preserve the high-level "
        "MATLAB/Python linear MPC formulation while moving toward a more realistic MuJoCo-integrated quadruped stack.",
    )

    heading(doc, "1. Background", level=1)
    paragraph(
        doc,
        "The previous Python baseline focused on reproducing the high-level structure of the MATLAB convex MPC: "
        "gait scheduling, reference rollout, nominal foothold generation, linearized single-rigid-body prediction, "
        "QP construction, and OSQP-based ground-reaction-force optimization. That baseline was useful for validating "
        "the MPC formulation itself, but it intentionally did not include a full low-level locomotion stack such as "
        "swing-leg inverse dynamics, touchdown logic, or torque-level realization.",
    )
    paragraph(
        doc,
        "Based on the feedback from the previous presentation, the current work treats rear-leg support, touchdown, "
        "and load transfer as the key bottleneck. For that reason, the controller was moved onto the stock "
        "Quadruped-PyMPC MuJoCo integration stack, while keeping only the custom linear_osqp controller at the "
        "high-level SRB force-planning layer.",
    )

    heading(doc, "2. Why the crawl scenario was used", level=1)
    paragraph(
        doc,
        "The crawl scenario was selected as a diagnostic scenario rather than a final performance benchmark. "
        "Compared with trot, crawl makes the contact transitions easier to interpret because legs move more "
        "sequentially. This is especially helpful for checking whether rear touchdown/recontact and load transfer "
        "are happening consistently.",
    )

    heading(doc, "3. Current same-scenario comparison", level=1)
    paragraph(
        doc,
        "To keep the comparison objective, both controllers were run under the same MuJoCo setting: "
        "robot=aliengo, scene=flat, gait=crawl, forward command vx=0.12, yaw=0.0. The stock controller used here is "
        "the sampling-based MPC provided by Quadruped-PyMPC, and the custom controller is the linear_osqp backend.",
    )

    add_picture(
        doc,
        CRAWL_DIR / "crawl_same_scenario_compare_table.png",
        width=6.9,
        caption_text="Figure 1. Same-scenario crawl comparison between stock sampling MPC and custom linear_osqp.",
    )

    paragraph(
        doc,
        "In this crawl setting, the stock sampling controller is not perfectly stable either, so it should not be "
        "treated as a perfect success case. However, it still provides a useful diagnostic reference because the "
        "custom linear_osqp backend shows lower mean body height and larger roll, indicating weaker robustness during "
        "contact transition.",
    )

    add_picture(
        doc,
        CRAWL_DIR / "stock_sampling_crawl_explainer.png",
        width=6.8,
        caption_text="Figure 2. Stock sampling MPC in the same crawl scenario.",
    )
    add_picture(
        doc,
        CRAWL_DIR / "linear_osqp_crawl_explainer.png",
        width=6.8,
        caption_text="Figure 3. Custom linear_osqp in the same crawl scenario.",
    )

    heading(doc, "4. Visual reference (GIF assets)", level=1)
    paragraph(
        doc,
        "The following GIF assets are embedded for quick qualitative inspection. Depending on the viewer, the GIFs may "
        "appear as static images inside the .docx file even though the original files are animated. The original GIF "
        "files are also saved next to this document in the outputs/report_progress_explainer/crawl_same_scenario folder.",
    )
    add_picture(
        doc,
        CRAWL_DIR / "stock_sampling_crawl.gif",
        width=5.4,
        caption_text="Figure 4. Stock sampling MPC crawl GIF (same scenario).",
    )
    add_picture(
        doc,
        CRAWL_DIR / "linear_osqp_crawl.gif",
        width=5.4,
        caption_text="Figure 5. Custom linear_osqp crawl GIF (same scenario).",
    )

    heading(doc, "5. Interpretation", level=1)
    paragraph(
        doc,
        "The current issue is not best described as a simple MPC formulation error. The SRB-level formulation itself "
        "is internally consistent, and the custom controller does run inside MuJoCo. The more accurate interpretation "
        "is that the present linear_osqp backend is still not robust enough at the interface between high-level force "
        "planning and low-level contact transition, particularly around rear touchdown/recontact.",
    )
    paragraph(
        doc,
        "It is also important to note that the current linear_osqp implementation is not just a bare linear QP. "
        "Additional contact-transition support logic has already been layered on top, including post-solve force "
        "regularization, support-force floors, rear-load biasing, pre-swing gating, late release logic, touchdown "
        "reacquire/confirm handling, and related bridge/recovery logic. These were not introduced as arbitrary tuning "
        "alone; they were added as structural heuristics inspired by the earlier prototype and by external legged "
        "locomotion references, then tuned within the current stack.",
    )

    heading(doc, "6. Next step", level=1)
    paragraph(
        doc,
        "The next step is to keep the stock integration stack fixed and continue improving the rear touchdown/recontact "
        "and post-touchdown stabilization behavior of the custom linear_osqp backend. In other words, the remaining "
        "work is focused less on changing the broad architecture and more on making the custom controller survive the "
        "critical contact-transition window more reliably.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "Appendix: Asset locations", level=1)
    paragraph(doc, f"DOCX output: {DOCX_PATH}")
    paragraph(doc, f"Crawl comparison asset folder: {CRAWL_DIR}")
    paragraph(doc, f"Table image: {CRAWL_DIR / 'crawl_same_scenario_compare_table.png'}")
    paragraph(doc, f"Stock explainer image: {CRAWL_DIR / 'stock_sampling_crawl_explainer.png'}")
    paragraph(doc, f"Linear explainer image: {CRAWL_DIR / 'linear_osqp_crawl_explainer.png'}")
    paragraph(doc, f"Stock GIF: {CRAWL_DIR / 'stock_sampling_crawl.gif'}")
    paragraph(doc, f"Linear GIF: {CRAWL_DIR / 'linear_osqp_crawl.gif'}")

    doc.save(str(DOCX_PATH))
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
