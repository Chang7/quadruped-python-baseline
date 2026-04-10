"""Generate a more polished weekly report v4 (Korean) for 2026-04-10."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WEEKLY_DIR = ROOT / "outputs" / "report_progress_explainer" / "weekly_progress_20260410"
COMPARE_DIR = ROOT / "outputs" / "report_progress_explainer" / "trot_stock_vs_custom_same_scenarios_20260410"
DOCX_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410_v4.docx"

SAME = ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios"
CRAWL_RELAX = ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck"

RUNS = {
    "straight_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_straight_default_20s" / "episode_000" / "summary.json",
    "turn_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_turn_default_10s" / "episode_000" / "summary.json",
    "disturb_before": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_baseline_before" / "episode_000" / "summary.json",
    "straight_stock": SAME / "stock_straight_20s" / "episode_000" / "summary.json",
    "straight_custom": SAME / "custom_straight_20s" / "episode_000" / "summary.json",
    "turn_stock": SAME / "stock_turn_10s" / "episode_000" / "summary.json",
    "turn_custom": SAME / "custom_turn_10s" / "episode_000" / "summary.json",
    "disturb_stock": SAME / "stock_disturb_4s" / "episode_000" / "summary.json",
    "disturb_custom": SAME / "custom_disturb_4s" / "episode_000" / "summary.json",
    "crawl_baseline": CRAWL_RELAX / "baseline_fy015_grf035" / "episode_000" / "summary.json",
    "crawl_fy100_grf035": CRAWL_RELAX / "fy100_grf035" / "episode_000" / "summary.json",
    "crawl_fy015_grf100": CRAWL_RELAX / "fy015_grf100" / "episode_000" / "summary.json",
    "crawl_fy100_grf100": CRAWL_RELAX / "fy100_grf100" / "episode_000" / "summary.json",
}

FIGURES = {
    "failure_ablation": WEEKLY_DIR / "weekly_failure_ablation_trot.png",
    "fyscale_recheck": WEEKLY_DIR / "weekly_trot_fyscale100_recheck.png",
    "tracking_overview": COMPARE_DIR / "trot_tracking_overview.png",
    "straight": COMPARE_DIR / "trot_straight_stock_vs_custom.png",
    "turn": COMPARE_DIR / "trot_turn_stock_vs_custom.png",
    "disturb": COMPARE_DIR / "trot_disturbance_stock_vs_custom.png",
    "crawl_relax": WEEKLY_DIR / "weekly_crawl_force_relax_recheck.png",
}


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt(v: float, d: int = 3) -> str:
    return f"{v:.{d}f}"


def _set_font(run, size: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], accent: str = "DCE6F1") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _shade_cell(cell, accent)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _set_font(r, 9, bold=True, color=RGBColor(31, 55, 86))
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            if r_idx % 2 == 1:
                _shade_cell(cell, "F8FBFE")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    _set_font(r, 9)


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, 9, color=RGBColor(90, 90, 90))
    run.italic = True


def _add_figure(doc: Document, path: Path, caption: str, width: float) -> None:
    if not path.exists():
        p = doc.add_paragraph(f"(누락된 그림) {path.name}")
        for r in p.runs:
            _set_font(r, 9, color=RGBColor(180, 50, 50))
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption(doc, caption)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_font(run, 10)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    _set_font(run, 13 if level == 1 else 11, bold=True, color=RGBColor(31, 55, 86))


def build() -> Path:
    data = {k: _load(v) for k, v in RUNS.items()}
    doc = Document()
    _set_doc_style(doc)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Weekly Progress Report")
    _set_font(run, 21, bold=True, color=RGBColor(31, 55, 86))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Quadruped-PyMPC MuJoCo integration + custom linear_osqp backend")
    _set_font(run, 11, color=RGBColor(90, 90, 90))
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026-04-10 weekly summary")
    _set_font(run, 10, color=RGBColor(120, 120, 120))

    doc.add_paragraph("")

    summary_box = doc.add_table(rows=1, cols=1)
    summary_box.style = "Table Grid"
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = summary_box.rows[0].cells[0]
    _shade_cell(cell, "F3F7FB")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "핵심 요약: 이번 주에는 새 제어기를 만든 것이 아니라, 기존 Python linear_osqp integration 경로 안에서 "
        "trot turn 성능을 제한하던 post-solve lateral-force clamp(fy_scale)를 병목으로 확인하고 제거했습니다. "
        "그 결과 trot straight / turn / disturbance 세 시나리오에서 roll·pitch 기준 posture quality가 stock reference 수준 이하가 되었습니다. "
        "반면 crawl에는 같은 접근이 통하지 않았고, 여전히 late seam / contact-transition 문제로 해석하는 것이 타당했습니다."
    )
    _set_font(run, 10)

    _heading(doc, "1. 이번 주 핵심 결과")
    doc.add_paragraph(
        "아래 표는 이번 주 시작 시점(before), 이번 주 최종 설정(after), 그리고 같은 local benchmark 조건에서 다시 확보한 stock reference를 정리한 것입니다."
    )
    posture_rows = [
        ["Straight 20 s mean|pitch|", _fmt(data["straight_before"]["mean_abs_pitch"]), _fmt(data["straight_custom"]["mean_abs_pitch"]), _fmt(data["straight_stock"]["mean_abs_pitch"])],
        ["Turn 10 s mean|roll|", _fmt(data["turn_before"]["mean_abs_roll"]), _fmt(data["turn_custom"]["mean_abs_roll"]), _fmt(data["turn_stock"]["mean_abs_roll"])],
        ["Turn 10 s mean|pitch|", _fmt(data["turn_before"]["mean_abs_pitch"]), _fmt(data["turn_custom"]["mean_abs_pitch"]), _fmt(data["turn_stock"]["mean_abs_pitch"])],
        ["Disturb 4 s mean|roll|", _fmt(data["disturb_before"]["mean_abs_roll"]), _fmt(data["disturb_custom"]["mean_abs_roll"]), _fmt(data["disturb_stock"]["mean_abs_roll"])],
        ["Disturb 4 s mean|pitch|", _fmt(data["disturb_before"]["mean_abs_pitch"]), _fmt(data["disturb_custom"]["mean_abs_pitch"]), _fmt(data["disturb_stock"]["mean_abs_pitch"])],
    ]
    _add_table(doc, ["Scenario", "Before", "After", "Stock"], posture_rows)
    _add_bullets(
        doc,
        [
            "현재 candidate setting에서는 straight 20 s, turn 10 s, disturbance 4 s가 모두 termination 없이 유지됩니다.",
            "roll / pitch 기준 posture quality는 세 시나리오 모두 stock reference 수준 이하입니다.",
            "다만 turn의 forward velocity tracking(mean_vx)은 여전히 stock이 더 높아, 전체 성능 우위라고 일반화하지는 않습니다.",
        ],
    )

    _heading(doc, "2. 시나리오 조건")
    doc.add_paragraph("이번 보고서에서의 trot 비교는 아래 exact local benchmark 조건으로 다시 실행한 same-scenario recheck를 기준으로 합니다.")
    scenario_rows = [
        ["straight", "20 s", "0.12 m/s", "0.0 rad/s", "없음"],
        ["turn", "10 s", "0.10 m/s", "0.3 rad/s", "없음"],
        ["disturbance", "4 s", "0.12 m/s", "0.0 rad/s", "x:0.5:0.25:4.0, x:2.3:0.25:8.0"],
    ]
    _add_table(doc, ["scenario", "duration", "speed", "yaw rate", "disturbance"], scenario_rows, accent="EAF2D3")
    doc.add_paragraph("공통 조건은 aliengo, flat ground, trot gait입니다.")

    _heading(doc, "3. 무엇을 확인했고 무엇이 실제로 효과가 있었는가")
    _add_bullets(
        doc,
        [
            "Q_theta_roll / Q_w_roll을 높이는 Q weighting 강화는 turn roll gap을 거의 줄이지 못했습니다.",
            "그 다음 단계에서, QP가 풀어낸 lateral force(fy)를 post-solve에서 다시 축소하는 fy_scale 구조가 남아 있음을 확인했습니다.",
            "dynamic_fy_roll_gain은 이 clamp를 완화하는 방향으로 부분 개선을 보였고, 최종적으로 fy_scale=1.0으로 clamp를 제거했을 때 가장 큰 개선이 나왔습니다.",
            "pitch는 random fluctuation보다 persistent bias 성격이 강했고, pitch_ref_offset=-0.03으로 상쇄가 가능했습니다.",
        ],
    )
    _add_figure(doc, FIGURES["failure_ablation"], "그림 1. 실패한 가설(Q weighting)과 효과가 있었던 방향(fy scaling 재검토)", 6.0)
    _add_figure(doc, FIGURES["fyscale_recheck"], "그림 2. fy_scale=1.0 재검증 결과", 6.0)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    _heading(doc, "4. 같은 시나리오에서 다시 본 stock vs custom")
    doc.add_paragraph(
        "아래 그림들은 same-scenario recheck 결과를 정리한 것입니다. straight는 XY path 정보량이 적어서 vx tracking을, "
        "turn / disturbance는 실제 경로 차이를 보기 위해 XY path를 사용했습니다."
    )
    _add_figure(doc, FIGURES["tracking_overview"], "그림 3. straight / turn / disturbance에서의 tracking / XY path 비교", 6.6)
    _add_figure(doc, FIGURES["straight"], "그림 4. straight 20 s: stock(좌) vs custom(우)", 6.4)
    _add_figure(doc, FIGURES["turn"], "그림 5. turn 10 s: stock(좌) vs custom(우)", 6.4)
    _add_figure(doc, FIGURES["disturb"], "그림 6. disturbance 4 s: stock(좌) vs custom(우)", 6.4)

    _heading(doc, "5. crawl은 왜 아직 별개 문제인가")
    doc.add_paragraph(
        "trot에서 효과가 컸던 force-authority 완화를 crawl에도 적용해 보았지만, fy_scale / grf_max_scale을 푼 네 가지 조합 모두 baseline보다 짧았습니다. "
        "즉 crawl은 단순한 force clamp 문제가 아니라 late rear load-transfer / post-touchdown stabilization seam 문제에 더 가깝습니다."
    )
    crawl_rows = [
        ["baseline fy=0.15, grf=0.35", _fmt(data["crawl_baseline"]["duration_s"]), _fmt(data["crawl_baseline"]["mean_abs_roll"]), _fmt(data["crawl_baseline"]["mean_abs_pitch"]), _fmt(data["crawl_baseline"]["mean_base_z"])],
        ["fy=1.0, grf=0.35", _fmt(data["crawl_fy100_grf035"]["duration_s"]), _fmt(data["crawl_fy100_grf035"]["mean_abs_roll"]), _fmt(data["crawl_fy100_grf035"]["mean_abs_pitch"]), _fmt(data["crawl_fy100_grf035"]["mean_base_z"])],
        ["fy=0.15, grf=1.0", _fmt(data["crawl_fy015_grf100"]["duration_s"]), _fmt(data["crawl_fy015_grf100"]["mean_abs_roll"]), _fmt(data["crawl_fy015_grf100"]["mean_abs_pitch"]), _fmt(data["crawl_fy015_grf100"]["mean_base_z"])],
        ["fy=1.0, grf=1.0", _fmt(data["crawl_fy100_grf100"]["duration_s"]), _fmt(data["crawl_fy100_grf100"]["mean_abs_roll"]), _fmt(data["crawl_fy100_grf100"]["mean_abs_pitch"]), _fmt(data["crawl_fy100_grf100"]["mean_base_z"])],
    ]
    _add_table(doc, ["crawl setting", "duration", "mean|roll|", "mean|pitch|", "mean base z"], crawl_rows, accent="FCE4D6")
    _add_figure(doc, FIGURES["crawl_relax"], "그림 7. crawl force-relax 재검증: 네 조합 모두 baseline보다 악화", 6.0)
    _add_bullets(
        doc,
        [
            "현재 custom crawl은 약 13.5 s까지 유지되지만 body height가 크게 낮아 functional locomotion보다 diagnostic scenario로 보는 편이 맞습니다.",
            "stock crawl도 같은 local setting에서 매우 짧게 종료되므로, crawl은 clean benchmark보다는 seam debugging scenario로 해석해야 합니다.",
        ],
    )

    _heading(doc, "6. 결론과 다음 단계")
    _add_bullets(
        doc,
        [
            "이번 주의 핵심 성과는 새 MPC 식 구현이 아니라, 기존 integration 경로 안의 post-solve scaling 병목을 찾아내고 제거한 것입니다.",
            "trot에서는 force realization path를 다시 보는 것이 cost weighting을 더 키우는 것보다 효과적이었습니다.",
            "crawl에서는 같은 접근이 통하지 않았으므로, 다음 단계는 force authority보다 seam state / transition timing을 더 직접 다루는 쪽이 자연스럽습니다.",
            "남은 확인 항목은 turn mean_vx gap, dynamic_fy_roll_gain의 필요성, 그리고 더 긴 horizon / 더 강한 disturbance에서의 재검증입니다.",
        ],
    )

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")
    return DOCX_PATH


if __name__ == "__main__":
    build()
