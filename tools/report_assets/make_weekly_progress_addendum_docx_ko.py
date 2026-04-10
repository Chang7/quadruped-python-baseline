from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
WEEKLY_DIR = ROOT / "outputs" / "report_progress_explainer" / "weekly_progress_20260410"
DOCX_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410_addendum.docx"
TROT_GRAPH_PATH = WEEKLY_DIR / "weekly_trot_fyscale100_recheck.png"
CRAWL_GRAPH_PATH = WEEKLY_DIR / "weekly_crawl_force_relax_recheck.png"

RUNS = {
    "turn_old": ROOT / "outputs" / "curated_runs" / "current" / "trot_turn_10s_g025_pitchoff003" / "episode_000" / "summary.json",
    "turn_new": WEEKLY_DIR / "custom_runs" / "trot_turn_10s_fyscale100" / "episode_000" / "summary.json",
    "turn_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_trot_turn_10s_weeklyref" / "episode_000" / "summary.json",
    "disturb_old": ROOT / "outputs" / "curated_runs" / "current" / "trot_disturb_4s_g025_pitchoff003" / "episode_000" / "summary.json",
    "disturb_new": WEEKLY_DIR / "custom_runs" / "trot_disturb_4s_fyscale100" / "episode_000" / "summary.json",
    "disturb_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_trot_disturb_4s_x48_recheck" / "episode_000" / "summary.json",
    "straight_new": WEEKLY_DIR / "custom_runs" / "trot_straight_4s_fyscale100" / "episode_000" / "summary.json",
    "straight_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_trot_4s_s012_isolated_recheck" / "episode_000" / "summary.json",
    "crawl_current": ROOT / "outputs" / "curated_runs" / "current" / "crawl_current_default_20s" / "episode_000" / "summary.json",
    "crawl_base": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "baseline_fy015_grf035" / "episode_000" / "summary.json",
    "crawl_fy100_grf035": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy100_grf035" / "episode_000" / "summary.json",
    "crawl_fy015_grf100": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy015_grf100" / "episode_000" / "summary.json",
    "crawl_fy100_grf100": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy100_grf100" / "episode_000" / "summary.json",
}

GIFS = {
    "straight": WEEKLY_DIR / "clips" / "trot_straight_custom_fyscale100.gif",
    "turn": WEEKLY_DIR / "clips" / "trot_turn_custom_fyscale100.gif",
    "disturb": WEEKLY_DIR / "clips" / "trot_disturb_custom_fyscale100.gif",
}


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _set_font(run) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_font(run)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_font(run)


def _picture(doc: Document, path: Path, width: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    _set_font(r)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption)
    cr.italic = True
    _set_font(cr)


def _style_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_font(run)


def build_doc() -> Path:
    data = {name: _load(path) for name, path in RUNS.items()}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("이번 주 진행 보고 추가 메모: fy_scale 재검증과 crawl force-relax 결과")
    r.bold = True
    r.font.size = Pt(17)
    _set_font(r)

    _paragraph(
        doc,
        "기존 주간 보고서 작성 이후, generic trot에서 post-solve lateral-force scaling(fy_scale)을 제거했을 때의 효과와 "
        "같은 접근이 crawl에도 통하는지를 추가로 재검증했다. 이번 추가 확인의 목적은 "
        "'trot에서 보인 개선이 단순한 튜닝 우연인지, 아니면 병목을 정확히 짚은 결과인지'를 빠르게 검증하는 것이었다.",
    )

    _heading(doc, "1. generic trot fy_scale=1.0 추가 확인", level=1)
    trot_table = doc.add_table(rows=1, cols=4)
    trot_table.style = "Table Grid"
    trot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = trot_table.rows[0].cells
    hdr[0].text = "Scenario"
    hdr[1].text = "기존 current"
    hdr[2].text = "fy_scale=1.0"
    hdr[3].text = "Stock"
    rows = [
        ("Turn 10 s mean|roll|", f"{data['turn_old']['mean_abs_roll']:.3f}", f"{data['turn_new']['mean_abs_roll']:.3f}", f"{data['turn_stock']['mean_abs_roll']:.3f}"),
        ("Turn 10 s mean|pitch|", f"{data['turn_old']['mean_abs_pitch']:.3f}", f"{data['turn_new']['mean_abs_pitch']:.3f}", f"{data['turn_stock']['mean_abs_pitch']:.3f}"),
        ("Disturb 4 s mean|roll|", f"{data['disturb_old']['mean_abs_roll']:.3f}", f"{data['disturb_new']['mean_abs_roll']:.3f}", f"{data['disturb_stock']['mean_abs_roll']:.3f}"),
        ("Disturb 4 s mean|pitch|", f"{data['disturb_old']['mean_abs_pitch']:.3f}", f"{data['disturb_new']['mean_abs_pitch']:.3f}", f"{data['disturb_stock']['mean_abs_pitch']:.3f}"),
        ("Straight 4 s mean|roll|", "-", f"{data['straight_new']['mean_abs_roll']:.3f}", f"{data['straight_stock']['mean_abs_roll']:.3f}"),
        ("Straight 4 s mean|pitch|", "-", f"{data['straight_new']['mean_abs_pitch']:.3f}", f"{data['straight_stock']['mean_abs_pitch']:.3f}"),
    ]
    for row in rows:
        cells = trot_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    _style_table(trot_table)

    _paragraph(
        doc,
        f"generic trot에서는 fy_scale=1.0이 안정적으로 동작했고, "
        f"turn 10 s mean|roll|은 {data['turn_old']['mean_abs_roll']:.3f}에서 {data['turn_new']['mean_abs_roll']:.3f}로, "
        f"disturb 4 s mean|roll|은 {data['disturb_old']['mean_abs_roll']:.3f}에서 {data['disturb_new']['mean_abs_roll']:.3f}로 줄었다. "
        "즉 trot에서는 기존 fy_scale clamp가 실제 성능 병목이었다는 해석이 훨씬 강해졌다.",
    )

    _picture(doc, TROT_GRAPH_PATH, 6.7, "Figure 1. generic trot에서 fy_scale=1.0 추가 재검증 결과")

    _heading(doc, "2. crawl force-relax 재검증", level=1)
    crawl_table = doc.add_table(rows=1, cols=5)
    crawl_table.style = "Table Grid"
    crawl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = crawl_table.rows[0].cells
    hdr[0].text = "crawl 20 s setting"
    hdr[1].text = "duration"
    hdr[2].text = "mean|roll|"
    hdr[3].text = "mean|pitch|"
    hdr[4].text = "mean base z"
    crawl_rows = [
        ("baseline fy=0.15, grf=0.35", data["crawl_base"]["duration_s"], data["crawl_base"]["mean_abs_roll"], data["crawl_base"]["mean_abs_pitch"], data["crawl_base"]["mean_base_z"]),
        ("fy=1.0, grf=0.35", data["crawl_fy100_grf035"]["duration_s"], data["crawl_fy100_grf035"]["mean_abs_roll"], data["crawl_fy100_grf035"]["mean_abs_pitch"], data["crawl_fy100_grf035"]["mean_base_z"]),
        ("fy=0.15, grf=1.0", data["crawl_fy015_grf100"]["duration_s"], data["crawl_fy015_grf100"]["mean_abs_roll"], data["crawl_fy015_grf100"]["mean_abs_pitch"], data["crawl_fy015_grf100"]["mean_base_z"]),
        ("fy=1.0, grf=1.0", data["crawl_fy100_grf100"]["duration_s"], data["crawl_fy100_grf100"]["mean_abs_roll"], data["crawl_fy100_grf100"]["mean_abs_pitch"], data["crawl_fy100_grf100"]["mean_base_z"]),
    ]
    for row in crawl_rows:
        cells = crawl_table.add_row().cells
        cells[0].text = row[0]
        for idx, value in enumerate(row[1:], start=1):
            cells[idx].text = f"{value:.3f}"
    _style_table(crawl_table)

    _paragraph(
        doc,
        f"crawl에서는 같은 방향이 통하지 않았다. current baseline과 같은 보수적 설정(fy=0.15, grf=0.35)은 {data['crawl_base']['duration_s']:.3f}s까지 유지됐지만, "
        f"fy만 풀면 {data['crawl_fy100_grf035']['duration_s']:.3f}s, grf만 풀면 {data['crawl_fy015_grf100']['duration_s']:.3f}s, "
        f"둘 다 풀면 {data['crawl_fy100_grf100']['duration_s']:.3f}s로 오히려 더 빨리 종료됐다.",
    )
    _bullet(doc, "즉 trot에서는 force clamp가 핵심 병목이었지만, crawl에서는 그렇지 않았다.")
    _bullet(doc, "crawl은 여전히 late rear load-transfer / post-touchdown seam 문제로 보는 해석이 더 타당하다.")
    _bullet(doc, "force authority를 넓히는 것만으로는 crawl의 failure location이 뒤로 밀리지 않았다.")

    _picture(doc, CRAWL_GRAPH_PATH, 6.7, "Figure 2. crawl에서 fy_scale / grf_max_scale 완화 재검증 결과")

    _heading(doc, "3. 추가 해석", level=1)
    _bullet(doc, "이번 추가 실험은 지난주와 이번 주 초의 ablation 결과를 더 선명하게 정리해준다.")
    _bullet(doc, "trot: Q weighting 강화보다 post-solve lateral-force realization이 더 중요했다.")
    _bullet(doc, "crawl: force limit 완화보다 contact-transition seam을 분리해서 보는 접근이 여전히 더 중요하다.")
    _bullet(doc, "따라서 다음 단계는 trot current를 fy_scale=1.0 기준으로 승격하고, crawl은 seam 로그/상태 분리를 계속 파는 것이 자연스럽다.")

    _heading(doc, "4. GIF 및 결과 위치", level=1)
    _bullet(doc, f"Straight GIF: {GIFS['straight']}")
    _bullet(doc, f"Turn GIF: {GIFS['turn']}")
    _bullet(doc, f"Disturb GIF: {GIFS['disturb']}")
    _bullet(doc, f"Trot fy_scale=1.0 graph: {TROT_GRAPH_PATH}")
    _bullet(doc, f"Crawl force-relax graph: {CRAWL_GRAPH_PATH}")
    _bullet(doc, f"Crawl force-relax runs: {ROOT / 'outputs' / 'archive' / 'raw_runs' / '20260410_crawl_force_relax_recheck'}")

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main() -> None:
    out = build_doc()
    print(out)


if __name__ == "__main__":
    main()
