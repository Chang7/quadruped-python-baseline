"""Generate the unified weekly progress report (Korean) for 2026-04-10."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
WEEKLY_DIR = ROOT / "outputs" / "report_progress_explainer" / "weekly_progress_20260410"
COMPARE_DIR = ROOT / "outputs" / "report_progress_explainer" / "trot_stock_vs_custom_same_scenarios_20260410"
DOCX_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410_unified.docx"
DOCX_FALLBACK_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410_unified_v2.docx"

RUNS = {
    "straight_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_straight_default_20s" / "episode_000" / "summary.json",
    "turn_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_turn_default_10s" / "episode_000" / "summary.json",
    "disturb_before": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_baseline_before" / "episode_000" / "summary.json",
    "straight_after": ROOT / "outputs" / "archive" / "raw_runs" / "trot_straight_20s_fyscale100" / "episode_000" / "summary.json",
    "turn_after": ROOT / "outputs" / "archive" / "raw_runs" / "trot_turn_10s_fyscale100" / "episode_000" / "summary.json",
    "disturb_after": ROOT / "outputs" / "archive" / "raw_runs" / "trot_disturb_4s_fyscale100" / "episode_000" / "summary.json",
    "straight_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_trot_straight_20s_weeklyref" / "episode_000" / "summary.json",
    "turn_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_trot_turn_10s_weeklyref" / "episode_000" / "summary.json",
    "disturb_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_trot_disturb_4s_x48_recheck" / "episode_000" / "summary.json",
    "turn_qsym": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_turn_10s_symroll_test" / "episode_000" / "summary.json",
    "disturb_qsym": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_symroll_test" / "episode_000" / "summary.json",
    "crawl_current": ROOT / "outputs" / "curated_runs" / "current" / "crawl_current_default_20s" / "episode_000" / "summary.json",
    "crawl_weakleg": ROOT / "outputs" / "archive" / "raw_runs" / "crawl_20260409" / "crawl_weakleg_share_ref040_test" / "episode_000" / "summary.json",
    "crawl_relax_baseline": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "baseline_fy015_grf035" / "episode_000" / "summary.json",
    "crawl_relax_fy100_grf035": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy100_grf035" / "episode_000" / "summary.json",
    "crawl_relax_fy015_grf100": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy015_grf100" / "episode_000" / "summary.json",
    "crawl_relax_fy100_grf100": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck" / "fy100_grf100" / "episode_000" / "summary.json",
    "stock_crawl_s003": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s003_isolated_recheck" / "episode_000" / "summary.json",
    "stock_crawl_s006": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s006_isolated_recheck" / "episode_000" / "summary.json",
    "stock_crawl_s012": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s012_isolated_recheck" / "episode_000" / "summary.json",
    "same_straight_stock": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "stock_straight_20s" / "episode_000" / "summary.json",
    "same_straight_custom": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "custom_straight_20s" / "episode_000" / "summary.json",
    "same_turn_stock": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "stock_turn_10s" / "episode_000" / "summary.json",
    "same_turn_custom": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "custom_turn_10s" / "episode_000" / "summary.json",
    "same_disturb_stock": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "stock_disturb_4s" / "episode_000" / "summary.json",
    "same_disturb_custom": ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios" / "custom_disturb_4s" / "episode_000" / "summary.json",
}

FIGURES = {
    "trot_fyscale100": WEEKLY_DIR / "weekly_trot_fyscale100_recheck.png",
    "trot_failure": WEEKLY_DIR / "weekly_failure_ablation_trot.png",
    "crawl_story": WEEKLY_DIR / "weekly_crawl_failure_story.png",
    "crawl_force_relax": WEEKLY_DIR / "weekly_crawl_force_relax_recheck.png",
    "same_summary": COMPARE_DIR / "trot_stock_vs_custom_summary_table.png",
    "same_straight": COMPARE_DIR / "trot_straight_stock_vs_custom.png",
    "same_turn": COMPARE_DIR / "trot_turn_stock_vs_custom.png",
    "same_disturb": COMPARE_DIR / "trot_disturbance_stock_vs_custom.png",
}

GIFS = {
    "trot_straight": WEEKLY_DIR / "clips" / "trot_straight_custom_fyscale100.gif",
    "trot_turn": WEEKLY_DIR / "clips" / "trot_turn_custom_fyscale100.gif",
    "trot_disturb": WEEKLY_DIR / "clips" / "trot_disturb_custom_fyscale100.gif",
    "crawl_mujoco": WEEKLY_DIR / "clips" / "crawl_custom_current_13p54s_mujoco.gif",
}


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _add_figure(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        doc.add_paragraph(f"(누락된 그림) {path}")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)


def _stock_crawl_range(data: dict[str, dict]) -> str:
    values = [
        data["stock_crawl_s003"]["duration_s"],
        data["stock_crawl_s006"]["duration_s"],
        data["stock_crawl_s012"]["duration_s"],
    ]
    return f"{min(values):.3f}~{max(values):.3f} s"


def build() -> Path:
    data = {name: _load(path) for name, path in RUNS.items() if path.exists()}
    doc = Document()
    _set_base_style(doc)

    title = doc.add_heading("이번 주 진행 보고", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Quadruped-PyMPC MuJoCo integration + custom linear_osqp backend")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 지난주 대비 이번 주 변화", level=1)
    doc.add_paragraph(
        "지난주에는 custom linear_osqp 경로가 trot에서 short-horizon viability는 보였지만 stock sampling 대비 "
        "posture deficit이 남아 있었고, crawl은 약 2.4 s에서 종료되었습니다. 당시 해석은 MPC 식 자체보다 "
        "planned force가 contact transition을 거쳐 실제로 구현되는 과정이 더 큰 병목이라는 것이었습니다."
    )
    doc.add_paragraph(
        "이번 주에는 그 가설을 시나리오별로 나눠 검증했습니다. 결과적으로 trot에서는 post-solve lateral-force "
        "축소 구조(fy_scale)가 핵심 병목이었고, 이를 제거한 뒤 straight / turn / disturbance 세 시나리오에서 "
        "roll·pitch 기준 posture quality가 stock reference 수준 이하로 내려갔습니다. 반면 crawl은 같은 방향의 "
        "force-authority 완화가 전혀 먹히지 않았고, 여전히 late rear load-transfer / post-touchdown seam 문제로 "
        "해석하는 것이 타당했습니다."
    )

    doc.add_heading("2. 핵심 수치", level=1)
    doc.add_paragraph(
        "아래 표는 이번 주 시작 시점(before), 이번 주 최종 설정(after), 그리고 stock reference를 같은 horizon 기준으로 "
        "정리한 값입니다. 이번 주 후 값은 fy_scale=1.0, dynamic_fy_roll_gain=0.25, pitch_ref_offset=-0.03 기준입니다."
    )
    summary_rows = []
    for label, scenario, metric in [
        ("Straight 20 s mean|pitch|", "straight", "mean_abs_pitch"),
        ("Turn 10 s mean|roll|", "turn", "mean_abs_roll"),
        ("Turn 10 s mean|pitch|", "turn", "mean_abs_pitch"),
        ("Disturb 4 s mean|roll|", "disturb", "mean_abs_roll"),
        ("Disturb 4 s mean|pitch|", "disturb", "mean_abs_pitch"),
    ]:
        summary_rows.append(
            [
                label,
                _fmt(data[f"{scenario}_before"][metric]),
                _fmt(data[f"{scenario}_after"][metric]),
                _fmt(data[f"{scenario}_stock"][metric]),
            ]
        )
    _add_table(doc, ["Scenario", "이번 주 전", "이번 주 후", "Stock"], summary_rows)
    doc.add_paragraph(
        "현재 candidate setting에서는 straight 20 s, turn 10 s, disturbance 4 s가 모두 termination 없이 유지됩니다. "
        "다만 이것이 전체 성능 우위라는 뜻은 아니며, forward velocity tracking 같은 다른 metric은 아직 stock과 차이가 남아 있습니다."
    )
    doc.add_paragraph(
        "이번 문서에서의 trot 비교는 아래와 같은 exact local benchmark 조건으로 다시 실행한 same-scenario recheck를 기준으로 정리했습니다."
    )
    scenario_rows = [
        ["straight", "20 s", "0.12 m/s", "0.0 rad/s", "없음"],
        ["turn", "10 s", "0.10 m/s", "0.3 rad/s", "없음"],
        ["disturbance", "4 s", "0.12 m/s", "0.0 rad/s", "x:0.5:0.25:4.0, x:2.3:0.25:8.0"],
    ]
    _add_table(doc, ["scenario", "duration", "speed", "yaw rate", "disturbance"], scenario_rows)
    doc.add_paragraph(
        "공통 조건은 aliengo, flat ground, trot gait입니다. 따라서 이 표에 들어간 stock/custom 값은 같은 시나리오 family일 뿐 아니라, "
        "이번 주에 동일한 local command 조건으로 다시 확보한 비교 결과입니다."
    )

    doc.add_heading("3. 이번 주에 실제로 한 것", level=1)
    for item in [
        "같은 시나리오를 유지한 상태에서 가설별 simulation sweep을 반복했습니다.",
        "새로운 QP 수식, 새로운 동역학 모델, 새로운 feedback path를 넣은 것은 없습니다.",
        "이번 주의 직접적인 코드 변경은 기존 Python integration 경로 안의 파라미터와 post-solve scaling 구조를 확인하고 조정한 것입니다.",
        "즉 이번 주 성과는 새로운 제어기 구현보다, 기존 구조 안에서 어떤 병목이 실제로 남아 있는지 분리해낸 데에 가깝습니다.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. 어떻게 찾았는가", level=1)
    for item in [
        "Q_theta_roll / Q_w_roll 대칭화: turn roll gap이 Q weighting 부족 때문인지 먼저 확인했습니다.",
        "dynamic_fy_roll_gain sweep: Q를 올려도 효과가 없던 이유가 실제 lateral-force realization 쪽에 있는지 확인했습니다.",
        "fy_scale 제거 실험: 부분 개선 이후, 아예 post-solve lateral-force clamp를 없애도 안정성이 유지되는지 확인했습니다.",
        "crawl force-relax 재검증: trot에서 먹힌 force-authority 완화가 crawl에도 통하는지 별도로 확인했습니다.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "즉 이번 주는 단순히 knob를 여러 개 돌린 것이 아니라, 지난주에 남은 질문을 가설로 바꾸고 같은 시나리오에서 하나씩 잘라서 확인한 주였습니다."
    )

    doc.add_heading("5. 무엇이 맞지 않았고, 무엇이 실제로 효과가 있었는가", level=1)
    doc.add_paragraph(
        "첫째, roll gap이 단순한 MPC weighting 문제라는 가설은 맞지 않았습니다. Q_theta_roll / Q_w_roll을 pitch 쪽과 대칭으로 "
        f"높여도 turn mean|roll|은 {_fmt(data['turn_before']['mean_abs_roll'])}에서 줄지 않았고, disturb 4 s에서도 "
        f"mean|roll|은 {_fmt(data['disturb_before']['mean_abs_roll'])} -> {_fmt(data['disturb_qsym']['mean_abs_roll'])} 수준의 "
        "소폭 변동에 그쳤습니다."
    )
    doc.add_paragraph(
        "둘째, crawl에서 넓은 weak-leg boost를 주면 더 버틸 것이라는 가설도 맞지 않았습니다. current default가 13.540 s까지 가는 반면, "
        f"weak_leg_share_ref=0.40 실험은 {_fmt(data['crawl_weakleg']['duration_s'])} s로 오히려 크게 regression했습니다."
    )
    doc.add_paragraph(
        "셋째, custom generic trot 경로 안에 남아 있던 post-solve fy_scale 구조가 실제 병목이라는 점이 이번 주에 확인되었습니다. "
        "dynamic_fy_roll_gain으로 clamp를 덜 거는 방향은 부분 개선을 보였고, 최종적으로 fy_scale=1.0으로 clamp를 제거했을 때 "
        "turn / disturbance posture가 가장 크게 좋아졌습니다."
    )
    doc.add_paragraph(
        "또한 pitch는 random fluctuation이라기보다 일관된 bias 성격이 강했고, pitch_ref_offset을 -0.03으로 조정했을 때 "
        "straight / turn / disturbance 전반에서 개선이 유지되었습니다."
    )
    _add_figure(doc, FIGURES["trot_fyscale100"], "그림 1. fy_scale=1.0 재검증 결과", width=6.0)
    _add_figure(doc, FIGURES["trot_failure"], "그림 2. 실패한 가설과 실제로 효과가 있었던 방향", width=6.0)

    doc.add_heading("6. 같은 시나리오에서 다시 본 stock vs custom 비교", level=1)
    doc.add_paragraph(
        "아래 비교는 straight / turn / disturbance 세 trot 시나리오를 같은 local setting으로 다시 돌려서 정리한 것입니다. "
        "즉 논문과 같은 시나리오 family(직진 / 회전 / disturbance)에는 대응하지만, 논문 command를 1:1로 재현했다는 뜻은 아닙니다."
    )
    same_rows = [
        [
            "straight 20 s",
            _fmt(data["same_straight_stock"]["mean_abs_roll"]),
            _fmt(data["same_straight_custom"]["mean_abs_roll"]),
            _fmt(data["same_straight_stock"]["mean_abs_pitch"]),
            _fmt(data["same_straight_custom"]["mean_abs_pitch"]),
            _fmt(data["same_straight_stock"]["mean_vx"]),
            _fmt(data["same_straight_custom"]["mean_vx"]),
        ],
        [
            "turn 10 s",
            _fmt(data["same_turn_stock"]["mean_abs_roll"]),
            _fmt(data["same_turn_custom"]["mean_abs_roll"]),
            _fmt(data["same_turn_stock"]["mean_abs_pitch"]),
            _fmt(data["same_turn_custom"]["mean_abs_pitch"]),
            _fmt(data["same_turn_stock"]["mean_vx"]),
            _fmt(data["same_turn_custom"]["mean_vx"]),
        ],
        [
            "disturb 4 s",
            _fmt(data["same_disturb_stock"]["mean_abs_roll"]),
            _fmt(data["same_disturb_custom"]["mean_abs_roll"]),
            _fmt(data["same_disturb_stock"]["mean_abs_pitch"]),
            _fmt(data["same_disturb_custom"]["mean_abs_pitch"]),
            _fmt(data["same_disturb_stock"]["mean_vx"]),
            _fmt(data["same_disturb_custom"]["mean_vx"]),
        ],
    ]
    _add_table(
        doc,
        ["scenario", "stock |roll|", "custom |roll|", "stock |pitch|", "custom |pitch|", "stock vx", "custom vx"],
        same_rows,
    )
    doc.add_paragraph(
        "이 same-scenario recheck에서는 roll / pitch 기준으로 custom이 세 시나리오 모두 stock보다 더 낮았습니다. "
        "다만 turn의 forward velocity tracking은 여전히 stock이 더 좋기 때문에, 전체 성능 우위로 일반화하기보다는 "
        "posture quality가 현재 custom 쪽에서 더 좋아졌다고 해석하는 것이 안전합니다."
    )
    _add_figure(doc, FIGURES["same_summary"], "그림 3. matched local trot benchmarks 요약 표", width=6.2)
    _add_figure(doc, FIGURES["same_straight"], "그림 4. trot + straight 같은 시나리오에서의 stock vs custom", width=6.4)
    _add_figure(doc, FIGURES["same_turn"], "그림 5. trot + turn 같은 시나리오에서의 stock vs custom", width=6.4)
    _add_figure(doc, FIGURES["same_disturb"], "그림 6. trot + disturbance 같은 시나리오에서의 stock vs custom", width=6.4)

    doc.add_heading("7. crawl은 여전히 다른 문제다", level=1)
    doc.add_paragraph(
        f"현재 custom crawl default는 {_fmt(data['crawl_current']['duration_s'])} s까지 유지되지만, mean base z가 "
        f"{_fmt(data['crawl_current']['mean_base_z'])} 수준으로 낮고 마지막 late seam에서 실패합니다. "
        f"반면 stock crawl은 같은 local setting에서 {_stock_crawl_range(data)} 정도로 매우 일찍 종료됩니다."
    )
    doc.add_paragraph(
        "중요한 점은, trot에서 효과가 컸던 force-authority 완화가 crawl에는 전혀 통하지 않았다는 것입니다. "
        "fy_scale / grf_max_scale을 풀어본 force-relax 재검증은 네 가지 조합 모두 baseline 13.540 s보다 나빴습니다."
    )
    doc.add_paragraph(
        "따라서 crawl은 functional locomotion benchmark라기보다 contact-transition diagnostic으로 보는 편이 맞고, "
        "현재 남은 병목은 force authority보다 rear close-handoff / late load-share / post-touchdown stabilization seam 쪽에 있습니다."
    )
    crawl_rows = [
        ["baseline fy=0.15, grf=0.35", _fmt(data["crawl_relax_baseline"]["duration_s"]), _fmt(data["crawl_relax_baseline"]["mean_abs_roll"]), _fmt(data["crawl_relax_baseline"]["mean_abs_pitch"]), _fmt(data["crawl_relax_baseline"]["mean_base_z"])],
        ["fy=1.0, grf=0.35", _fmt(data["crawl_relax_fy100_grf035"]["duration_s"]), _fmt(data["crawl_relax_fy100_grf035"]["mean_abs_roll"]), _fmt(data["crawl_relax_fy100_grf035"]["mean_abs_pitch"]), _fmt(data["crawl_relax_fy100_grf035"]["mean_base_z"])],
        ["fy=0.15, grf=1.0", _fmt(data["crawl_relax_fy015_grf100"]["duration_s"]), _fmt(data["crawl_relax_fy015_grf100"]["mean_abs_roll"]), _fmt(data["crawl_relax_fy015_grf100"]["mean_abs_pitch"]), _fmt(data["crawl_relax_fy015_grf100"]["mean_base_z"])],
        ["fy=1.0, grf=1.0", _fmt(data["crawl_relax_fy100_grf100"]["duration_s"]), _fmt(data["crawl_relax_fy100_grf100"]["mean_abs_roll"]), _fmt(data["crawl_relax_fy100_grf100"]["mean_abs_pitch"]), _fmt(data["crawl_relax_fy100_grf100"]["mean_base_z"])],
    ]
    _add_table(doc, ["crawl 20 s setting", "duration", "mean|roll|", "mean|pitch|", "mean base z"], crawl_rows)
    _add_figure(doc, FIGURES["crawl_story"], "그림 7. crawl current default와 weak-leg failure 비교", width=6.0)
    _add_figure(doc, FIGURES["crawl_force_relax"], "그림 8. crawl force-relax 재검증 결과", width=6.0)

    doc.add_heading("8. 정성적 자료(GIF / MuJoCo visual)", level=1)
    doc.add_paragraph(
        "숫자와 그래프 외에 실제 움직임은 아래 GIF로 바로 확인할 수 있습니다. 특히 turn / disturbance는 posture quality 개선이 시각적으로도 분명합니다."
    )
    for label, path in [
        ("trot straight custom", GIFS["trot_straight"]),
        ("trot turn custom", GIFS["trot_turn"]),
        ("trot disturbance custom", GIFS["trot_disturb"]),
        ("crawl custom 13.54 s MuJoCo", GIFS["crawl_mujoco"]),
    ]:
        doc.add_paragraph(f"{label}: {path}", style="List Bullet")

    doc.add_heading("9. 남은 문제와 다음 단계", level=1)
    for item in [
        "crawl 13.540 s late-seam failure는 여전히 미해결이며, force authority가 아니라 seam state / transition timing 쪽을 계속 파야 합니다.",
        "fy_scale=1.0 상태에서 dynamic_fy_roll_gain=0.25가 아직 필요한지는 별도 재검증할 가치가 있습니다.",
        "trot는 posture quality는 좋아졌지만, turn case의 mean_vx처럼 velocity-tracking metric은 아직 더 확인이 필요합니다.",
        "즉 다음 단계는 무작정 더 많은 gain tuning보다, crawl seam 로직과 force-realization path의 구조를 다시 보는 쪽이 자연스럽습니다.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10. 이번 결과가 시사하는 점", level=1)
    doc.add_paragraph(
        "이번 주 결과는 'MPC cost를 더 키우는 것'보다 '이미 풀어낸 답이 실제로 어떻게 적용되는지'를 점검하는 것이 더 중요할 수 있다는 점을 보여줍니다. "
        "QP 내부 friction cone 제약 자체가 lateral force를 막고 있었던 것이 아니라, Python integration 경로 안의 post-solve scaling이 "
        "trot turn 성능의 핵심 병목이었다는 점이 확인되었습니다."
    )
    doc.add_paragraph(
        "반대로 crawl에서는 같은 방식의 force-authority 완화가 전혀 통하지 않았기 때문에, 남은 문제는 단순한 힘 부족이 아니라 "
        "contact transition / post-touchdown stabilization seam에 더 가깝다는 해석이 강화되었습니다."
    )

    doc.add_heading("부록: 사용한 결과물 위치", level=1)
    for path in [
        WEEKLY_DIR,
        COMPARE_DIR,
        ROOT / "outputs" / "archive" / "raw_runs" / "20260410_trot_stock_vs_custom_same_scenarios",
        ROOT / "outputs" / "archive" / "raw_runs" / "20260410_crawl_force_relax_recheck",
    ]:
        doc.add_paragraph(str(path), style="List Bullet")

    try:
        doc.save(str(DOCX_PATH))
        saved = DOCX_PATH
    except PermissionError:
        doc.save(str(DOCX_FALLBACK_PATH))
        saved = DOCX_FALLBACK_PATH

    print(f"Saved: {saved}")
    return saved


if __name__ == "__main__":
    build()
