from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
WEEKLY_DIR = ROOT / "outputs" / "report_progress_explainer" / "weekly_progress_20260410"
DOCX_PATH = WEEKLY_DIR / "weekly_progress_report_ko_20260410.docx"

RUNS = {
    "trot_straight_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_straight_default_20s" / "episode_000" / "summary.json",
    "trot_straight_after": ROOT / "outputs" / "curated_runs" / "current" / "trot_straight_20s_g025_pitchoff003" / "episode_000" / "summary.json",
    "trot_straight_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_trot_straight_20s_weeklyref" / "episode_000" / "summary.json",
    "trot_turn_before": ROOT / "outputs" / "curated_runs" / "predecessors" / "trot_current_turn_default_10s" / "episode_000" / "summary.json",
    "trot_turn_after": ROOT / "outputs" / "curated_runs" / "current" / "trot_turn_10s_g025_pitchoff003" / "episode_000" / "summary.json",
    "trot_turn_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_trot_turn_10s_weeklyref" / "episode_000" / "summary.json",
    "trot_disturb_before": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_baseline_before" / "episode_000" / "summary.json",
    "trot_disturb_after": ROOT / "outputs" / "curated_runs" / "current" / "trot_disturb_4s_g025_pitchoff003" / "episode_000" / "summary.json",
    "trot_disturb_stock": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_trot_disturb_4s_x48_recheck" / "episode_000" / "summary.json",
    "turn_symroll_fail": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_turn_10s_symroll_test" / "episode_000" / "summary.json",
    "disturb_symroll_fail": ROOT / "outputs" / "archive" / "raw_runs" / "trot_20260409" / "quality_sweeps" / "trot_disturb_4s_symroll_test" / "episode_000" / "summary.json",
    "crawl_current": ROOT / "outputs" / "curated_runs" / "current" / "crawl_current_default_20s" / "episode_000" / "summary.json",
    "crawl_weakleg_fail": ROOT / "outputs" / "archive" / "raw_runs" / "crawl_20260409" / "crawl_weakleg_share_ref040_test" / "episode_000" / "summary.json",
    "crawl_stock_s003": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s003_isolated_recheck" / "episode_000" / "summary.json",
    "crawl_stock_s006": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s006_isolated_recheck" / "episode_000" / "summary.json",
    "crawl_stock_s012": ROOT / "outputs" / "curated_runs" / "stock_baselines" / "stock_sampling_crawl_4s_s012_isolated_recheck" / "episode_000" / "summary.json",
}

FIGURES = {
    "trot_summary": WEEKLY_DIR / "weekly_trot_improvement_summary.png",
    "trot_failure": WEEKLY_DIR / "weekly_failure_ablation_trot.png",
    "crawl_story": WEEKLY_DIR / "weekly_crawl_failure_story.png",
    "stock_vs_custom": WEEKLY_DIR / "weekly_stock_vs_custom.png",
}

CLIPS = {
    "trot_straight": WEEKLY_DIR / "clips" / "trot_straight_current.gif",
    "trot_turn": WEEKLY_DIR / "clips" / "trot_turn_current.gif",
    "trot_disturb": WEEKLY_DIR / "clips" / "trot_disturb_current.gif",
    "crawl_current": WEEKLY_DIR / "clips" / "crawl_current.gif",
}


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _loads() -> dict[str, dict]:
    return {name: _load(path) for name, path in RUNS.items()}


def set_normal_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_picture(doc: Document, path: Path, width: float, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption(doc, caption_text)


def add_core_metrics_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Scenario"
    hdr[1].text = "이번 주 전"
    hdr[2].text = "이번 주 후"
    hdr[3].text = "Stock"
    rows = [
        ("Straight 20s mean|pitch|", "0.052", "0.031", "0.045"),
        ("Turn 10s mean|roll|", "0.024", "0.018", "0.014"),
        ("Turn 10s mean|pitch|", "0.059", "0.043", "0.045"),
        ("Disturb 4s mean|roll|", "0.024", "0.020", "0.023"),
        ("Disturb 4s mean|pitch|", "0.053", "0.034", "0.055"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_doc() -> Path:
    data = _loads()
    stock_crawl_min = min(
        data["crawl_stock_s003"]["duration_s"],
        data["crawl_stock_s006"]["duration_s"],
        data["crawl_stock_s012"]["duration_s"],
    )
    stock_crawl_max = max(
        data["crawl_stock_s003"]["duration_s"],
        data["crawl_stock_s006"]["duration_s"],
        data["crawl_stock_s012"]["duration_s"],
    )

    doc = Document()
    set_normal_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_page_number(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("이번 주 진행 보고")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Quadruped-PyMPC MuJoCo integration + custom linear_osqp backend")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    heading(doc, "1. 핵심 수치", level=1)
    paragraph(
        doc,
        "이번 주에는 새로운 제어 식을 추가하기보다, 기존 제어 구조 안에서 matched scenario sweep을 반복해 "
        "어떤 파라미터 조정이 실제 품질 개선으로 이어지는지 확인했습니다. 아래 표는 이번 주 전후의 핵심 수치만 모은 것입니다.",
    )
    add_core_metrics_table(doc)
    paragraph(
        doc,
        "이번 주 후 candidate setting에서는 straight 20s, turn 10s, disturb 4s trot를 모두 termination 없이 통과했고, "
        "특히 pitch는 straight, turn, disturbance에서 모두 줄었습니다.",
    )
    bullet(
        doc,
        f"또한 stock crawl도 현재 local setting에서 4s를 완주하지 못했고, 세 reference run이 "
        f"{stock_crawl_min:.3f}s ~ {stock_crawl_max:.3f}s에서 모두 종료된다는 점을 다시 확인했습니다. "
        "따라서 crawl은 main performance benchmark라기보다 contact-transition diagnostic scenario로 보는 것이 맞습니다.",
    )

    heading(doc, "2. 이번 주에 한 것", level=1)
    bullet(doc, "기존 제어 구조를 유지한 채 matched scenario sweep으로 남아 있던 질문을 검증했습니다.")
    bullet(doc, "새로운 제어 모듈, QP 구조 변경, 물리 모델 변경은 없었습니다.")
    bullet(doc, "stock sampling 경로와 stock 코드는 건드리지 않고 custom linear_osqp 프로필 파라미터만 조정했습니다.")
    bullet(doc, "실제로 바꾼 핵심 파라미터는 dynamic_fy_roll_gain (0 -> 0.25)과 pitch_ref_offset (-0.01 -> -0.03)입니다.")
    bullet(doc, "weak_leg_share_ref는 0.40까지 시도했지만 regression을 확인한 뒤 다시 0.0으로 원복했습니다.")
    bullet(doc, "crawl은 성능 benchmark가 아니라 rear touchdown/recontact와 late load-transfer seam을 보는 diagnostic scenario로 계속 사용했습니다.")

    heading(doc, "3. 틀린 가설", level=1)
    bullet(
        doc,
        "Q_theta_roll / Q_w_roll을 pitch 쪽과 대칭으로 강화하면 turn roll gap이 줄어들 것이라는 가설은 맞지 않았습니다. "
        f"실제로 turn mean|roll|은 {data['trot_turn_before']['mean_abs_roll']:.3f}에서 "
        f"{data['turn_symroll_fail']['mean_abs_roll']:.3f}으로 사실상 변하지 않았고, "
        f"disturb 4s mean|roll|도 {data['trot_disturb_before']['mean_abs_roll']:.3f}에서 "
        f"{data['disturb_symroll_fail']['mean_abs_roll']:.3f}으로 소폭 변동하는 데 그쳤습니다. "
        "즉 Q weighting 강화만으로는 roll gap이 해소되지 않았습니다.",
    )
    bullet(
        doc,
        "crawl에서 weak_leg_share_ref=0.40 같은 aggressive weak-leg boost를 주면 안정성이 좋아질 것이라는 가설도 틀렸습니다. "
        f"current default duration {data['crawl_current']['duration_s']:.3f}s 대비 weak-leg test는 "
        f"{data['crawl_weakleg_fail']['duration_s']:.3f}s로 크게 regression했고, 정상 crawl alternation 구간에서도 과발동하는 문제가 확인됐습니다.",
    )
    bullet(
        doc,
        "stock crawl을 clean gold baseline으로 보는 해석도 적절하지 않았습니다. "
        f"현재 local setting에서 stock crawl은 세 기준 run 모두 {stock_crawl_min:.3f}s ~ {stock_crawl_max:.3f}s에서 종료되므로, "
        "crawl은 성능 비교보다 contact-transition diagnostic으로 해석하는 편이 더 자연스럽습니다.",
    )

    heading(doc, "4. 맞았던 가설", level=1)
    bullet(
        doc,
        "turn roll gap은 MPC cost weighting 강화보다 existing lateral-force-related gain 조정이 더 효과적이었습니다. "
        f"turn mean|roll|은 {data['trot_turn_before']['mean_abs_roll']:.3f}에서 "
        f"{data['trot_turn_after']['mean_abs_roll']:.3f}으로 개선됐습니다.",
    )
    bullet(
        doc,
        "pitch는 random fluctuation보다 persistent bias 성격이 강했고, reference offset으로 상쇄할 수 있었습니다. "
        f"straight 20s mean|pitch|은 {data['trot_straight_before']['mean_abs_pitch']:.3f}에서 "
        f"{data['trot_straight_after']['mean_abs_pitch']:.3f}으로 줄었습니다.",
    )
    bullet(
        doc,
        "crawl은 broad recovery를 한 덩어리로 키우는 것보다 seam을 분리해서 보는 접근이 더 생산적이었습니다. "
        "이제 failure location은 초기 rear recontact seam이 아니라 마지막 low-height late seam으로 좁혀졌습니다.",
    )

    heading(doc, "5. 남은 문제", level=1)
    bullet(
        doc,
        f"turn roll은 여전히 stock reference ({data['trot_turn_stock']['mean_abs_roll']:.3f}) 대비 current custom "
        f"({data['trot_turn_after']['mean_abs_roll']:.3f}) 사이에 gap이 남아 있습니다.",
    )
    bullet(
        doc,
        f"crawl current default는 {data['crawl_current']['duration_s']:.3f}s까지 가지만 마지막 late seam에서 invalid contact로 종료되며, 아직 완전한 해결은 아닙니다.",
    )
    bullet(
        doc,
        f"straight 20s mean_vx는 {data['trot_straight_before']['mean_vx']:.3f}에서 "
        f"{data['trot_straight_after']['mean_vx']:.3f}으로 소폭 낮아졌습니다.",
    )

    heading(doc, "6. 해석과 다음 단계", level=1)
    paragraph(
        doc,
        "이번 주 결과는 단순히 MPC cost를 더 키우는 것보다, 실제 보행에서 corrective force가 어떻게 구현되고 "
        "contact transition이 언제 어떻게 전환되는지가 더 중요하다는 점을 보여줍니다. 다만 이 결론은 이번 주의 "
        "파라미터 sweep 결과에 기반한 해석이고, 다음 단계에서 구조적 변경이 필요할 경우 어떤 방향을 먼저 볼지에 대한 힌트를 준 것으로 이해하는 것이 적절합니다.",
    )
    paragraph(
        doc,
        "이번 주의 개선은 본질적으로 기존 구조 안에서의 튜닝이었습니다. 즉 새로운 제어 식이나 새로운 feedback path를 구현한 것이 아니라, "
        "이미 정의되어 있던 knob의 영향을 분리해서 본 것입니다. 특히 QP 내부 cost weighting 강화는 기대한 효과를 거의 내지 못했고, "
        "QP 입력 reference 조정이나 QP 이후 force shaping 쪽 조정이 실제 품질 개선으로 이어졌습니다.",
    )
    paragraph(
        doc,
        "현재 generic profile의 lateral force는 QP가 풀어낸 해를 사후에 fy_scale로 스케일링하는 구조이며, "
        "dynamic_fy_roll_gain도 그 스케일을 조정하는 성격에 가깝습니다. pitch_ref_offset 역시 persistent bias를 "
        "reference로 보상하는 우회입니다. 따라서 이번 주의 sweep 결과는 post-hoc 조정만으로 갈 수 있는 범위를 상당 부분 확인한 것이고, "
        "다음 단계에서는 QP 내부 lateral-force cost 구조, 모델 기반 feedforward, 혹은 구조적 bias 보정 같은 방향을 검토하는 것이 자연스럽습니다.",
    )

    heading(doc, "7. 그래프 결과", level=1)
    add_picture(doc, FIGURES["trot_summary"], 6.8, "그림 1. 이번 주 trot 핵심 수치 개선 요약")
    add_picture(doc, FIGURES["trot_failure"], 6.8, "그림 2. 실패한 가설과 실제로 효과 있었던 방향")
    add_picture(doc, FIGURES["crawl_story"], 6.8, "그림 3. crawl current default와 weak-leg failure 비교")
    add_picture(doc, FIGURES["stock_vs_custom"], 6.8, "그림 4. same-horizon stock vs custom 비교")

    heading(doc, "8. GIF 및 시각 자료 위치", level=1)
    paragraph(
        doc,
        "GIF는 본문 표와 그래프를 보조하는 시각 증거로 첨부하는 것이 좋습니다. Word 환경에 따라 GIF가 정지 이미지처럼 보일 수 있으므로, "
        "문서 안에는 그래프와 표를 넣고 GIF는 같은 폴더의 첨부 파일로 함께 전달하는 방식을 권장합니다.",
    )
    bullet(doc, f"그래프 및 보고 폴더: {WEEKLY_DIR}")
    bullet(doc, f"straight GIF: {CLIPS['trot_straight']}")
    bullet(doc, f"turn GIF: {CLIPS['trot_turn']}")
    bullet(doc, f"disturb GIF: {CLIPS['trot_disturb']}")
    bullet(doc, f"crawl GIF: {CLIPS['crawl_current']}")

    heading(doc, "9. 참고 run 위치", level=1)
    bullet(doc, f"turn before: {RUNS['trot_turn_before'].parent}")
    bullet(doc, f"turn after: {RUNS['trot_turn_after'].parent}")
    bullet(doc, f"straight before: {RUNS['trot_straight_before'].parent}")
    bullet(doc, f"straight after: {RUNS['trot_straight_after'].parent}")
    bullet(doc, f"disturb before: {RUNS['trot_disturb_before'].parent}")
    bullet(doc, f"disturb after: {RUNS['trot_disturb_after'].parent}")
    bullet(doc, f"crawl current: {RUNS['crawl_current'].parent}")
    bullet(doc, f"crawl weak-leg fail: {RUNS['crawl_weakleg_fail'].parent}")
    bullet(doc, f"stock crawl 0.03: {RUNS['crawl_stock_s003'].parent}")
    bullet(doc, f"stock crawl 0.06: {RUNS['crawl_stock_s006'].parent}")
    bullet(doc, f"stock crawl 0.12: {RUNS['crawl_stock_s012'].parent}")

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "부록: 수치 근거", level=1)
    paragraph(
        doc,
        "이 문서에서 사용한 핵심 summary 값은 straight 20s, turn 10s, disturb 4s, crawl 20s run의 summary.json에서 읽은 값입니다. "
        "stock reference는 가능한 한 같은 horizon 기준으로 다시 정리한 값만 사용했습니다.",
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(DOCX_PATH))
        return DOCX_PATH
    except PermissionError:
        fallback = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_new{DOCX_PATH.suffix}")
        doc.save(str(fallback))
        return fallback


if __name__ == "__main__":
    print(build_doc())
